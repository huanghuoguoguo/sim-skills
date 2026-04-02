from __future__ import annotations

import os
import tempfile
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Length

# Import from sibling module using relative path resolution
import sys
libs_path = Path(__file__).parent
if str(libs_path) not in sys.path:
    sys.path.insert(0, str(libs_path))

from docx_parser_models import ParagraphFact, StyleFact, WordDocumentFacts, HeaderFooterFact


def parse_word_document(path: str | Path) -> WordDocumentFacts:
    source_path = Path(path)
    open_path, cleanup = normalize_wordprocessingml_package(source_path)
    try:
        document = Document(str(open_path))
        paragraphs = [
            build_paragraph_fact(index, paragraph)
            for index, paragraph in enumerate(document.paragraphs)
        ]
        styles = collect_styles(document)
        headers = extract_headers_footers(document, "header")
        footers = extract_headers_footers(document, "footer")
        return WordDocumentFacts(
            format=normalized_format(source_path),
            metadata={
                "filename": source_path.name,
                "title": read_core_property(document, "title"),
                "paragraph_count": len(document.paragraphs),
                "non_empty_paragraph_count": sum(1 for paragraph in paragraphs if paragraph.text),
            },
            layout=extract_layout(document),
            paragraphs=paragraphs,
            styles=styles,
            headers=headers,
            footers=footers,
        )
    finally:
        cleanup()


def build_paragraph_fact(index: int, paragraph) -> ParagraphFact:
    style_name = paragraph.style.name if paragraph.style is not None else None
    style_id = paragraph.style.style_id if paragraph.style is not None else None
    properties, property_sources = extract_paragraph_properties(paragraph)
    return ParagraphFact(
        id=f"p-{index + 1}",
        index=index,
        text=paragraph.text.strip(),
        style_name=style_name,
        style_id=style_id,
        properties=properties,
        property_sources=property_sources,
    )


def collect_styles(document: Document) -> list[StyleFact]:
    styles: list[StyleFact] = []
    for style in document.styles:
        if getattr(style, "type", None) is None:
            continue
        styles.append(
            StyleFact(
                name=style.name,
                style_id=getattr(style, "style_id", None),
                style_type=str(style.type),
                base_style=getattr(getattr(style, "base_style", None), "name", None),
                properties=extract_style_properties(style),
            )
        )
    return styles


def normalized_format(path: Path) -> str:
    suffix = path.suffix.lower().lstrip(".")
    return suffix or "docx"


def normalize_wordprocessingml_package(path: Path) -> tuple[Path, callable]:
    if path.suffix.lower() not in {".dotm", ".docm"}:
        return path, lambda: None

    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_path = Path(temp_file.name)
    temp_file.close()

    with ZipFile(path) as src, ZipFile(temp_path, "w", compression=ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "[Content_Types].xml":
                data = rewrite_main_document_content_type(data)
            dst.writestr(info, data)

    return temp_path, lambda: safe_unlink(temp_path)


def rewrite_main_document_content_type(data: bytes) -> bytes:
    root = ET.fromstring(data)
    namespace = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    for override in root.findall("ct:Override", namespace):
        if override.attrib.get("PartName") == "/word/document.xml":
            override.set(
                "ContentType",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
            )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def safe_unlink(path: Path) -> None:
    try:
        os.unlink(path)
    except FileNotFoundError:
        return


def read_core_property(document: Document, name: str) -> str | None:
    core_properties = getattr(document, "core_properties", None)
    if core_properties is None:
        return None
    value = getattr(core_properties, name, None)
    return value or None


def extract_layout(document: Document) -> dict[str, Any]:
    section = document.sections[0] if document.sections else None
    if section is None:
        return {}
    return {
        "page_size": {
            "width_cm": length_to_cm(section.page_width),
            "height_cm": length_to_cm(section.page_height),
        },
        "page_margins": {
            "top_cm": length_to_cm(section.top_margin),
            "bottom_cm": length_to_cm(section.bottom_margin),
            "left_cm": length_to_cm(section.left_margin),
            "right_cm": length_to_cm(section.right_margin),
        },
        "header_distance_cm": length_to_cm(section.header_distance),
        "footer_distance_cm": length_to_cm(section.footer_distance),
    }


def extract_headers_footers(document: Document, hf_type: str) -> list[dict]:
    """提取页眉或页脚内容。

    Args:
        document: Document 对象
        hf_type: "header" 或 "footer"

    Returns:
        HeaderFooterFact 列表
    """
    result = []
    for i, section in enumerate(document.sections):
        # 获取页眉或页脚 - try primary attribute first
        hf = getattr(section, f"{hf_type}_part", None)
        if hf is None:
            # Fallback to internal attribute (some python-docx versions)
            hf = getattr(section, f"_{hf_type}", None)

        if hf is None:
            continue

        paragraphs = []
        texts = []
        hf_paragraphs = getattr(hf, "paragraphs", [])
        for j, para in enumerate(hf_paragraphs):
            para_text = getattr(para, "text", "")
            if para_text:
                para_text = para_text.strip()
            if para_text:
                texts.append(para_text)
            paragraphs.append({
                "index": j,
                "text": para_text,
                "style_name": getattr(getattr(para, "style", None), "name", None),
            })

        if paragraphs or texts:
            result.append({
                "type": hf_type,
                "section_index": i,
                "paragraphs": paragraphs,
                "text": " | ".join(texts),
            })

    return result


def extract_paragraph_properties(paragraph) -> tuple[dict[str, Any], dict[str, str]]:
    properties: dict[str, Any] = {}
    property_sources: dict[str, str] = {}
    fmt = paragraph.paragraph_format
    style_properties = extract_style_properties(paragraph.style)

    assign_property(properties, property_sources, "alignment",
                    normalize_alignment(first_present(paragraph.alignment,
                                   style_properties.get("alignment"))),
                    "direct" if paragraph.alignment is not None else "style")

    line_spacing_mode, line_spacing_value = normalize_line_spacing(
        first_present(fmt.line_spacing_rule, style_properties.get("line_spacing_rule")),
        first_present(fmt.line_spacing, style_properties.get("line_spacing")),
    )
    assign_property(properties, property_sources, "line_spacing_mode", line_spacing_mode,
                    "direct" if fmt.line_spacing_rule is not None or fmt.line_spacing is not None else "style")
    assign_property(properties, property_sources, "line_spacing_value", line_spacing_value,
                    "direct" if fmt.line_spacing is not None else "style")

    assign_property(properties, property_sources, "space_before_pt",
                    length_to_pt(first_present(fmt.space_before, style_properties.get("space_before"))),
                    "direct" if fmt.space_before is not None else "style")
    assign_property(properties, property_sources, "space_after_pt",
                    length_to_pt(first_present(fmt.space_after, style_properties.get("space_after"))),
                    "direct" if fmt.space_after is not None else "style")

    # 中英文字体分离提取
    run_fonts = detect_font_from_runs_detailed(paragraph)
    style_fonts = extract_style_fonts_detailed(paragraph.style)

    # 中文字体 (eastAsia)
    font_family_east_asia = first_present(run_fonts.get("font_family_east_asia"), style_fonts.get("font_family_east_asia"), style_properties.get("font_family"))
    assign_property(properties, property_sources, "font_family_east_asia", font_family_east_asia, "direct" if run_fonts.get("font_family_east_asia") else "style")

    # 英文字体 (ascii)
    font_family_ascii = first_present(run_fonts.get("font_family_ascii"), style_fonts.get("font_family_ascii"), style_properties.get("font_family"))
    assign_property(properties, property_sources, "font_family_ascii", font_family_ascii, "direct" if run_fonts.get("font_family_ascii") else "style")

    # 同时保留单一 font_family，便于简单消费方读取。
    font_family = font_family_east_asia or font_family_ascii
    assign_property(properties, property_sources, "font_family", font_family, "direct" if font_family else "style")

    font_size_pt = first_present(detect_font_size_from_runs(paragraph), style_properties.get("font_size_pt"))
    assign_property(properties, property_sources, "font_size_pt", font_size_pt, "direct" if detect_font_size_from_runs(paragraph) else "style")

    first_line_indent_pt = length_to_pt(first_present(fmt.first_line_indent, style_properties.get("first_line_indent")))
    assign_property(properties, property_sources, "first_line_indent_pt", first_line_indent_pt, "direct" if fmt.first_line_indent is not None else "style")

    return properties, property_sources


def assign_property(
    properties: dict[str, Any],
    property_sources: dict[str, str],
    name: str,
    value: Any,
    source: str,
) -> None:
    if value is None:
        return
    properties[name] = value
    property_sources[name] = source


def extract_style_properties(style) -> dict[str, Any]:
    if style is None:
        return {}

    properties: dict[str, Any] = {}
    for current_style in style_chain(style):
        # 使用详细字体提取
        style_fonts = extract_style_fonts_detailed(current_style)
        if style_fonts["font_family_east_asia"]:
            properties["font_family_east_asia"] = style_fonts["font_family_east_asia"]
        if style_fonts["font_family_ascii"]:
            properties["font_family_ascii"] = style_fonts["font_family_ascii"]

        # 同时保留单一 font_family，便于简单消费方读取。
        style_font = getattr(current_style, "font", None)
        if style_font is not None:
            font_name = first_present(style_font.name, extract_style_font_name(current_style))
            if font_name is not None and "font_family" not in properties:
                properties["font_family"] = font_name
            if style_font.size is not None:
                properties["font_size_pt"] = round(style_font.size.pt, 2)

        paragraph_format = getattr(current_style, "paragraph_format", None)
        if paragraph_format is None:
            continue
        if paragraph_format.alignment is not None:
            properties["alignment"] = paragraph_format.alignment
        if paragraph_format.line_spacing is not None:
            properties["line_spacing"] = paragraph_format.line_spacing
        if paragraph_format.line_spacing_rule is not None:
            properties["line_spacing_rule"] = paragraph_format.line_spacing_rule
        if paragraph_format.space_before is not None:
            properties["space_before"] = paragraph_format.space_before
        if paragraph_format.space_after is not None:
            properties["space_after"] = paragraph_format.space_after
        if paragraph_format.first_line_indent is not None:
            properties["first_line_indent"] = paragraph_format.first_line_indent
    return properties


def style_chain(style) -> list[Any]:
    chain: list[Any] = []
    current = style
    while current is not None:
        chain.append(current)
        current = getattr(current, "base_style", None)
    return list(reversed(chain))


def extract_style_font_name(style) -> str | None:
    """Extract a single font name from style for simple consumers."""
    element = getattr(style, "_element", None)
    if element is None or getattr(element, "rPr", None) is None:
        return None
    fonts = element.rPr.rFonts
    if fonts is None:
        return None
    for attr in ("eastAsia", "ascii", "hAnsi"):
        value = fonts.get(qn(f"w:{attr}"))
        if value:
            return value
    return None


def extract_style_fonts_detailed(style) -> dict[str, str | None]:
    """Extract separate fonts for Chinese and Latin text from style."""
    result = {"font_family_east_asia": None, "font_family_ascii": None}

    element = getattr(style, "_element", None)
    if element is None or getattr(element, "rPr", None) is None:
        return result

    fonts = element.rPr.rFonts
    if fonts is not None:
        east_asia = fonts.get(qn(f"w:eastAsia"))
        ascii_font = fonts.get(qn(f"w:ascii"))
        if east_asia:
            result["font_family_east_asia"] = east_asia
        if ascii_font:
            result["font_family_ascii"] = ascii_font

    return result


def detect_font_from_runs(paragraph) -> str | None:
    """Detect a single font name from runs for simple consumers."""
    for run in paragraph.runs:
        font_name = extract_run_font_name(run)
        if font_name:
            return font_name
    return None


def detect_font_from_runs_detailed(paragraph) -> dict[str, str | None]:
    """Detect fonts from runs, returning separate fonts for Chinese and Latin text."""
    result = {"font_family_east_asia": None, "font_family_ascii": None}
    last_run = None

    for run in paragraph.runs:
        last_run = run
        r_pr = getattr(run._element, "rPr", None)
        if r_pr is None or r_pr.rFonts is None:
            continue

        fonts = r_pr.rFonts
        if not result["font_family_east_asia"]:
            value = fonts.get(qn("w:eastAsia"))
            if value:
                result["font_family_east_asia"] = value
        if not result["font_family_ascii"]:
            value = fonts.get(qn("w:ascii"))
            if value:
                result["font_family_ascii"] = value

        # Early exit if both fonts found
        if result["font_family_east_asia"] and result["font_family_ascii"]:
            break

    # Fallback to run.font.name if no explicit font found
    if last_run is not None:
        if not result["font_family_ascii"] and last_run.font.name:
            result["font_family_ascii"] = last_run.font.name
        if not result["font_family_east_asia"] and last_run.font.name:
            result["font_family_east_asia"] = last_run.font.name

    return result


def detect_font_size_from_runs(paragraph) -> float | None:
    for run in paragraph.runs:
        if run.font.size is not None:
            return round(run.font.size.pt, 2)
    return None


def extract_run_font_name(run) -> str | None:
    if run.font.name:
        return run.font.name
    r_pr = getattr(run._element, "rPr", None)
    if r_pr is None or r_pr.rFonts is None:
        return None
    for attr in ("eastAsia", "ascii", "hAnsi"):
        value = r_pr.rFonts.get(qn(f"w:{attr}"))
        if value:
            return value
    return None


def normalize_alignment(value: Any) -> str | None:
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
    }
    return mapping.get(value)


def normalize_line_spacing(rule: Any, value: Any) -> tuple[str | None, float | None]:
    if value is None and rule is None:
        return None, None
    rule_map = {
        WD_LINE_SPACING.SINGLE: "single",
        WD_LINE_SPACING.ONE_POINT_FIVE: "multiple",
        WD_LINE_SPACING.DOUBLE: "multiple",
        WD_LINE_SPACING.EXACTLY: "exact",
        WD_LINE_SPACING.AT_LEAST: "at_least",
        WD_LINE_SPACING.MULTIPLE: "multiple",
    }
    normalized_rule = rule_map.get(rule)
    if isinstance(value, Length):
        return normalized_rule, round(value.pt, 2)
    if isinstance(value, float):
        return normalized_rule, round(value, 2)
    return normalized_rule, None


def length_to_pt(value: Any) -> float | None:
    if value is None:
        return None
    return round(value.pt, 2)


def length_to_cm(value: Any) -> float | None:
    if value is None:
        return None
    return round(value.cm, 2)


def first_present(*values: Any) -> Any:
    for value in values:
        if value is not None:
            return value
    return None
